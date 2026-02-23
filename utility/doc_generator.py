import io
from docx import Document
from docx.shared import Inches, Pt
import re

def generate_docx(topic: str, markdown_content: str) -> io.BytesIO:
    """
    Parses a markdown string and generates an in-memory DOCX file.
    Returns the BytesIO stream containing the document.
    """
    document = Document()
    
    # Title
    title = document.add_heading(f"Deep Research Report: {topic}", 0)
    title.alignment = 1 # Center align

    # Very basic markdown parser for DOCX
    lines = markdown_content.split('\n')
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        # Headings
        if line_stripped.startswith('### '):
            document.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith('## '):
            document.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('# '):
            document.add_heading(line_stripped[2:], level=1)
        
        # Bullet points
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            document.add_paragraph(line_stripped[2:], style='List Bullet')
            
        # Numbered lists
        elif re.match(r'^\d+\.\s', line_stripped):
            # Extract text after the number and dot
            text = re.sub(r'^\d+\.\s+', '', line_stripped)
            document.add_paragraph(text, style='List Number')
            
        # Normal Text
        else:
            # Strip bold/italic markdown for simplicity in this MVP DOCX builder
            clean_text = line_stripped.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
            document.add_paragraph(clean_text)

    # Save to memory stream
    doc_stream = io.BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream
